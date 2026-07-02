#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Exportación final de la sesión a Word (.docx).

El documento contiene los JSON de todas las búsquedas de la sesión CONCATENADOS en
un único array JSON válido, en fuente monoespaciada, pensado para pasarlo después a
un lector de coincidencias. Antes del JSON se incluye un breve resumen de la sesión.

Deduplicación: por (fuente, jobid) — la misma vacante hallada en varias búsquedas
aparece UNA vez, con el campo extra "busquedas" listando qué búsquedas la produjeron
(mismo criterio que build_excel.py).

Requiere: pip install python-docx
"""
import json
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor


def consolidar(sesion: dict) -> list:
    """
    Une las vacantes de todas las búsquedas en una sola lista deduplicada.
    Cada vacante conserva sus campos originales y gana "busquedas": ["occ: analista / cdmx", ...].
    """
    orden = []       # claves en orden de aparición
    por_clave = {}   # clave -> vacante
    terminos = {}    # clave -> set de términos de búsqueda

    for b in sesion.get("busquedas", []):
        etiqueta = f"{b.get('sitio')}: {b.get('empleo')} / {b.get('ciudad')}"
        for v in b.get("vacantes", []):
            clave = (v.get("fuente"), v.get("jobid"))
            # Sin jobid no se puede deduplicar con seguridad: se trata como única.
            if clave[1] is None:
                clave = (v.get("fuente"), id(v))
            if clave not in por_clave:
                por_clave[clave] = dict(v)  # copia para no mutar la sesión
                orden.append(clave)
                terminos[clave] = []
            if etiqueta not in terminos[clave]:
                terminos[clave].append(etiqueta)

    resultado = []
    for clave in orden:
        v = por_clave[clave]
        v["busquedas"] = terminos[clave]
        resultado.append(v)
    return resultado


def exportar(sesion: dict, ruta_docx: str) -> int:
    """
    Genera el .docx con el JSON consolidado. Devuelve el número de vacantes únicas.
    """
    vacantes = consolidar(sesion)

    doc = Document()

    # --- Encabezado / resumen de la sesión ----------------------------------
    titulo = doc.add_heading("Vacantes recopiladas — sesión de búsqueda", level=1)

    p = doc.add_paragraph()
    r = p.add_run(
        f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}   |   "
        f"Búsquedas: {len(sesion.get('busquedas', []))}   |   "
        f"Vacantes únicas: {len(vacantes)}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    for b in sesion.get("busquedas", []):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(
            f"{b.get('sitio')} — «{b.get('empleo')}» en «{b.get('ciudad')}» "
            f"({len(b.get('vacantes', []))} vacantes, {b.get('timestamp', '')})")
        r.font.size = Pt(10)

    doc.add_heading("JSON consolidado", level=2)
    nota = doc.add_paragraph()
    r = nota.add_run(
        "El siguiente bloque es un array JSON válido (una vacante por bloque, "
        "deduplicada por fuente+jobid, con el campo \"busquedas\" añadido).")
    r.font.size = Pt(9)
    r.italic = True

    # --- JSON concatenado en monoespaciada -----------------------------------
    # Un párrafo por vacante (JSON indentado) para que Word no sufra con un
    # párrafo gigante y el documento siga siendo un array JSON válido al leerlo.
    _parrafo_mono(doc, "[")
    for i, v in enumerate(vacantes):
        texto = json.dumps(v, ensure_ascii=False, indent=2)
        if i < len(vacantes) - 1:
            texto += ","
        _parrafo_mono(doc, texto)
    _parrafo_mono(doc, "]")

    doc.save(ruta_docx)
    return len(vacantes)


def _parrafo_mono(doc, texto: str) -> None:
    """Añade un párrafo en fuente monoespaciada pequeña."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(texto)
    r.font.name = "Courier New"
    r.font.size = Pt(8)
